from __future__ import annotations

import os
import time
import uuid
from pathlib import Path

try:
    import pytesseract
except ModuleNotFoundError:
    pytesseract = None
    
from docx import Document
from docx.shared import Pt
from PIL import Image, ImageFilter, ImageOps, UnidentifiedImageError
from pillow_heif import register_heif_opener
from werkzeug.utils import secure_filename

from app.models.file_asset import FileAsset

register_heif_opener()


def ensure_ocr_docx_output_directory(path: str | Path | None = None) -> Path:
    directory = Path(path or os.getenv("DOCX_OUTPUT_DIR", "app/temp_storage/converted"))
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def build_ocr_docx_filename(source_files: list[FileAsset]) -> str:
    timestamp = int(time.time())
    token = uuid.uuid4().hex[:8]

    if len(source_files) == 1:
        base_name = secure_filename(source_files[0].original_filename) or "file"
        stem = Path(base_name).stem or "file"
        return f"{stem}-ocr-{timestamp}-{token}.docx"

    return f"merged-ocr-{timestamp}-{token}.docx"


def preprocess_image_for_ocr(source_path: str | Path) -> Image.Image:
    try:
        with Image.open(source_path) as image:
            grayscale = ImageOps.grayscale(image)
            contrasted = ImageOps.autocontrast(grayscale)
            sharpened = contrasted.filter(ImageFilter.SHARPEN)
            return sharpened
    except UnidentifiedImageError as exc:
        raise ValueError(f"Unsupported or unreadable image file for OCR: {source_path}") from exc


def normalize_extracted_text(text: str) -> str:
    normalized = "\n".join(line.strip() for line in text.splitlines())
    normalized = "\n".join(line for line in normalized.splitlines() if line)
    min_length = int(os.getenv("OCR_MIN_TEXT_LENGTH", "3"))

    if len(normalized.strip()) < min_length:
        return os.getenv("OCR_DOCX_EMPTY_TEXT_PLACEHOLDER", "[No readable text detected]")

    return normalized.strip()


def extract_text_from_image(source_path: str | Path, languages: str | None = None) -> str:
    languages = languages or os.getenv("OCR_LANGUAGES", "eng")
    
    if pytesseract is None:
        raise RuntimeError(
            "pytesseract is not installed. Install pytesseract and ensure Tesseract OCR is available."
    )

    try:
        processed_image = preprocess_image_for_ocr(source_path)
        text = pytesseract.image_to_string(processed_image, lang=languages)
        return normalize_extracted_text(text)
    except pytesseract.pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError(
            "Tesseract OCR is not installed on this system. Install 'tesseract-ocr' first."
        ) from exc


def create_editable_docx_from_images(
    source_files: list[FileAsset],
    title: str | None = None,
    output_dir: str | Path | None = None,
) -> dict:
    if not source_files:
        raise ValueError("At least one source file is required to create an OCR DOCX.")

    output_directory = ensure_ocr_docx_output_directory(output_dir)
    output_filename = build_ocr_docx_filename(source_files)
    output_path = output_directory / output_filename

    document = Document()
    document_title = title or "OCR Extracted Text"
    document.add_heading(document_title, level=1)

    include_source_names = os.getenv("OCR_DOCX_INCLUDE_SOURCE_NAMES", "true").lower() == "true"
    extracted_sections = 0

    for source_file in source_files:
        file_path = Path(source_file.storage_path)
        if not file_path.exists():
            raise FileNotFoundError(
                f"Uploaded source file is missing on disk: {source_file.original_filename}"
            )

        extracted_text = extract_text_from_image(file_path)

        if include_source_names:
            document.add_paragraph(source_file.original_filename)

        for paragraph_text in extracted_text.split("\n"):
            paragraph = document.add_paragraph(paragraph_text)
            paragraph.paragraph_format.space_after = Pt(6)

        extracted_sections += 1

    document.save(output_path)

    return {
        "output_filename": output_filename,
        "output_path": str(output_path),
        "output_format": "docx",
        "size_bytes": output_path.stat().st_size,
        "extracted_section_count": extracted_sections,
    }
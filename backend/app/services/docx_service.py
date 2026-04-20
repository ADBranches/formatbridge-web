from __future__ import annotations

import os
import time
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
from pillow_heif import register_heif_opener
from werkzeug.utils import secure_filename

from app.models.file_asset import FileAsset

register_heif_opener()


def ensure_docx_output_directory(path: str | Path | None = None) -> Path:
    directory = Path(path or os.getenv("DOCX_OUTPUT_DIR", "app/temp_storage/converted"))
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def build_docx_filename(source_files: list[FileAsset]) -> str:
    timestamp = int(time.time())
    token = uuid.uuid4().hex[:8]

    if len(source_files) == 1:
        base_name = secure_filename(source_files[0].original_filename) or "file"
        stem = Path(base_name).stem or "file"
        return f"{stem}-{timestamp}-{token}.docx"

    return f"merged-{timestamp}-{token}.docx"


def get_image_width_inches(image_path: str | Path, max_width_inches: float) -> float:
    with Image.open(image_path) as image:
        width_px, height_px = image.size
        if width_px <= 0 or height_px <= 0:
            return max_width_inches

        dpi = image.info.get("dpi", (96, 96))[0] or 96
        width_inches = width_px / dpi if dpi else width_px / 96

        if width_inches <= 0:
            return max_width_inches

        return min(width_inches, max_width_inches)


def create_docx_from_images(
    source_files: list[FileAsset],
    title: str | None = None,
    output_dir: str | Path | None = None,
) -> dict:
    if not source_files:
        raise ValueError("At least one source file is required to create a DOCX.")

    output_directory = ensure_docx_output_directory(output_dir)
    output_filename = build_docx_filename(source_files)
    output_path = output_directory / output_filename

    document = Document()

    include_title = os.getenv("DOCX_INCLUDE_TITLE", "true").lower() == "true"
    effective_title = title or os.getenv("DOCX_DEFAULT_TITLE", "Converted Images")

    if include_title:
        document.add_heading(effective_title, level=1)

    max_width_inches = float(os.getenv("DOCX_IMAGE_MAX_WIDTH_INCHES", "6.5"))
    spacing_after_pt = int(os.getenv("DOCX_IMAGE_SPACING_AFTER_PT", "12"))

    inserted_count = 0

    for source_file in source_files:
        file_path = Path(source_file.storage_path)
        if not file_path.exists():
            raise FileNotFoundError(
                f"Uploaded source file is missing on disk: {source_file.original_filename}"
            )

        document.add_paragraph(source_file.original_filename)

        width_inches = get_image_width_inches(file_path, max_width_inches)

        image_paragraph = document.add_paragraph()
        run = image_paragraph.add_run()
        run.add_picture(str(file_path), width=Inches(width_inches))
        image_paragraph.paragraph_format.space_after = Pt(spacing_after_pt)

        inserted_count += 1

    document.save(output_path)

    return {
        "output_filename": output_filename,
        "output_path": str(output_path),
        "output_format": "docx",
        "size_bytes": output_path.stat().st_size,
        "embedded_image_count": inserted_count,
    }